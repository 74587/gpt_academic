import os
import time
from abc import ABC, abstractmethod
from datetime import datetime
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import  Inches, Cm
from docx.shared import Pt, RGBColor, Inches
from typing import Dict, List, Tuple


class DocumentFormatter(ABC):
    """文档格式化基类，定义文档格式化的基本接口"""

    def __init__(self, final_summary: str, file_summaries_map: Dict, failed_files: List[Tuple]):
        self.final_summary = final_summary
        self.file_summaries_map = file_summaries_map
        self.failed_files = failed_files

    @abstractmethod
    def format_failed_files(self) -> str:
        """格式化失败文件列表"""
        pass

    @abstractmethod
    def format_file_summaries(self) -> str:
        """格式化文件总结内容"""
        pass

    @abstractmethod
    def create_document(self) -> str:
        """创建完整文档"""
        pass


class WordFormatter(DocumentFormatter):
    """Word格式文档生成器 - 符合中国政府公文格式规范(GB/T 9704-2012)，并进行了优化"""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.doc = Document()
        self._setup_document()
        self._create_styles()
        # 初始化标题编号系统 - 只使用两级编号
        self.numbers = {
            1: 0,  # 一级标题编号
            2: 0  # 二级标题编号
        }

    def _setup_document(self):
        """设置文档基本格式"""
        sections = self.doc.sections
        for section in sections:
            # 设置页面大小为A4
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # 设置页边距
            section.top_margin = Cm(3.7)  # 上边距37mm
            section.bottom_margin = Cm(3.5)  # 下边距35mm
            section.left_margin = Cm(2.8)  # 左边距28mm
            section.right_margin = Cm(2.6)  # 右边距26mm
            # 设置页眉页脚
            section.header_distance = Cm(2.0)
            section.footer_distance = Cm(2.0)

    def _create_styles(self):
        """创建文档样式"""
        # 创建正文样式
        style = self.doc.styles.add_style('Normal_Custom', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = '仿宋'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(14)  # 调整正文字号为14号
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Pt(28)  # 首行缩进两个字符（14pt * 2）

        # 创建各级标题样式（从大到小递减）
        self._create_heading_style('Title_Custom', '方正小标宋简体', 32, WD_PARAGRAPH_ALIGNMENT.CENTER)  # 大标题，增大字号到32
        self._create_heading_style('Heading1_Custom', '黑体', 22, WD_PARAGRAPH_ALIGNMENT.LEFT)  # 一级标题
        self._create_heading_style('Heading2_Custom', '黑体', 18, WD_PARAGRAPH_ALIGNMENT.LEFT)  # 二级标题
        self._create_heading_style('Heading3_Custom', '黑体', 16, WD_PARAGRAPH_ALIGNMENT.LEFT)  # 三级标题

    def _create_heading_style(self, style_name: str, font_name: str, font_size: int, alignment):
        """创建标题样式"""
        style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        style.font.size = Pt(font_size)
        style.font.bold = True  # 所有标题都加粗
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        return style

    def _get_heading_number(self, level: int) -> str:
        """生成标题编号"""
        if level == 0:  # 主标题不需要编号
            return ""

        self.numbers[level] += 1  # 增加当前级别的编号

        # 如果是一级标题，重置二级标题编号
        if level == 1:
            self.numbers[2] = 0

        # 根据级别返回不同格式的编号
        if level == 1:
            return f"{self.numbers[1]}. "
        elif level == 2:
            return f"{self.numbers[1]}.{self.numbers[2]} "
        return ""

    def _add_heading(self, text: str, level: int):
        """添加带编号的标题"""
        style_map = {
            0: 'Title_Custom',
            1: 'Heading1_Custom',
            2: 'Heading2_Custom',
            3: 'Heading3_Custom'
        }

        # 获取标题编号
        number = self._get_heading_number(level)

        # 创建段落
        paragraph = self.doc.add_paragraph(style=style_map[level])

        # 分别添加编号和文本，并设置样式
        if number:
            number_run = paragraph.add_run(number)
            self._get_run_style(number_run, '黑体', 22 if level == 1 else 18, True)

        text_run = paragraph.add_run(text)
        font_size = 32 if level == 0 else (22 if level == 1 else 18)  # 主标题32号，一级标题22号，其他18号
        self._get_run_style(text_run, '黑体', font_size, True)

        # 特殊处理：主标题添加日期
        if level == 0:
            date_paragraph = self.doc.add_paragraph()
            date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            date_run = date_paragraph.add_run(datetime.now().strftime('%Y年%m月%d日'))
            date_run.font.name = '仿宋'
            date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            date_run.font.size = Pt(16)

        return paragraph

    def _get_run_style(self, run, font_name: str, font_size: int, bold: bool = False):
        """设置文本运行对象的样式"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold

    def format_failed_files(self) -> str:
        """格式化失败文件列表"""
        result = []
        if not self.failed_files:
            return "\n".join(result)

        result.append("处理失败文件:")
        for fp, reason in self.failed_files:
            result.append(f"• {os.path.basename(fp)}: {reason}")

        # 在文档中添加内容
        self._add_heading("处理失败文件", 1)
        for fp, reason in self.failed_files:
            self._add_content(f"• {os.path.basename(fp)}: {reason}", indent=False)
        self.doc.add_paragraph()

        return "\n".join(result)

    def _add_content(self, text: str, indent: bool = True):
        """添加正文内容"""
        paragraph = self.doc.add_paragraph(text, style='Normal_Custom')
        if not indent:
            paragraph.paragraph_format.first_line_indent = Pt(0)  # 不缩进的段落
        return paragraph

    def format_file_summaries(self) -> str:
        """格式化文件总结内容"""
        result = []
        sorted_paths = sorted(self.file_summaries_map.keys())
        current_dir = ""

        for path in sorted_paths:
            dir_path = os.path.dirname(path)
            if dir_path != current_dir:
                if dir_path:
                    result.append(f"\n📁 {dir_path}")
                    self._add_heading(f"📁 {dir_path}", 2)
                current_dir = dir_path

            # 添加文件名和内容到结果字符串
            file_name = os.path.basename(path)
            result.append(f"\n📄 {file_name}")
            result.append(self.file_summaries_map[path])

            # 在文档中添加文件名作为带编号的二级标题
            self._add_heading(f"📄 {file_name}", 2)
            self._add_content(self.file_summaries_map[path])
            self.doc.add_paragraph()

        return "\n".join(result)

    def create_document(self):
        """创建完整Word文档并返回文档对象"""
        # 重置所有编号
        for level in self.numbers:
            self.numbers[level] = 0

        # 添加主标题（更大字号和加粗）
        self._add_heading("文档总结报告", 0)
        self.doc.add_paragraph()

        # 添加总体摘要
        self._add_heading("总体摘要", 1)
        self._add_content(self.final_summary)
        self.doc.add_paragraph()

        # 添加失败文件列表（如果有）
        if self.failed_files:
            self.format_failed_files()

        # 添加文件详细总结
        self._add_heading("各文件详细总结", 1)
        self.format_file_summaries()

        return self.doc  # 返回文档对象


class MarkdownFormatter(DocumentFormatter):
    """Markdown格式文档生成器"""

    def format_failed_files(self) -> str:
        if not self.failed_files:
            return ""

        formatted_text = ["\n## ⚠️ 处理失败的文件"]
        for fp, reason in self.failed_files:
            formatted_text.append(f"- {os.path.basename(fp)}: {reason}")
        formatted_text.append("\n---")
        return "\n".join(formatted_text)

    def format_file_summaries(self) -> str:
        formatted_text = []
        sorted_paths = sorted(self.file_summaries_map.keys())
        current_dir = ""

        for path in sorted_paths:
            dir_path = os.path.dirname(path)
            if dir_path != current_dir:
                if dir_path:
                    formatted_text.append(f"\n## 📁 {dir_path}")
                current_dir = dir_path

            file_name = os.path.basename(path)
            formatted_text.append(f"\n### 📄 {file_name}")
            formatted_text.append(self.file_summaries_map[path])
            formatted_text.append("\n---")

        return "\n".join(formatted_text)

    def create_document(self) -> str:
        document = [
            "# 📑 文档总结报告",
            "\n## 总体摘要",
            self.final_summary
        ]

        if self.failed_files:
            document.append(self.format_failed_files())

        document.extend([
            "\n# 📚 各文件详细总结",
            self.format_file_summaries()
        ])

        return "\n".join(document)


class HtmlFormatter(DocumentFormatter):
    """HTML格式文档生成器"""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.css_styles = """
        body {
            font-family: "Microsoft YaHei", Arial, sans-serif;
            line-height: 1.6;
            max-width: 1000px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }
        h1 {
            color: #2c3e50;
            border-bottom: 2px solid #eee;
            padding-bottom: 10px;
            font-size: 24px;
            text-align: center;
        }
        h2 {
            color: #34495e;
            margin-top: 30px;
            font-size: 20px;
            border-left: 4px solid #3498db;
            padding-left: 10px;
        }
        h3 {
            color: #2c3e50;
            font-size: 18px;
            margin-top: 20px;
        }
        .summary {
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 5px;
            margin: 20px 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .details {
            margin-top: 40px;
        }
        .failed-files {
            background-color: #fff3f3;
            padding: 15px;
            border-left: 4px solid #e74c3c;
            margin: 20px 0;
        }
        .file-summary {
            background-color: #fff;
            padding: 15px;
            margin: 15px 0;
            border-radius: 4px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }
        """

    def format_failed_files(self) -> str:
        if not self.failed_files:
            return ""

        failed_files_html = ['<div class="failed-files">']
        failed_files_html.append("<h2>⚠️ 处理失败的文件</h2>")
        failed_files_html.append("<ul>")
        for fp, reason in self.failed_files:
            failed_files_html.append(f"<li><strong>{os.path.basename(fp)}:</strong> {reason}</li>")
        failed_files_html.append("</ul></div>")
        return "\n".join(failed_files_html)

    def format_file_summaries(self) -> str:
        formatted_html = []
        sorted_paths = sorted(self.file_summaries_map.keys())
        current_dir = ""

        for path in sorted_paths:
            dir_path = os.path.dirname(path)
            if dir_path != current_dir:
                if dir_path:
                    formatted_html.append(f'<h2>📁 {dir_path}</h2>')
                current_dir = dir_path

            file_name = os.path.basename(path)
            formatted_html.append('<div class="file-summary">')
            formatted_html.append(f'<h3>📄 {file_name}</h3>')
            formatted_html.append(f'<p>{self.file_summaries_map[path]}</p>')
            formatted_html.append('</div>')

        return "\n".join(formatted_html)

    def create_document(self) -> str:
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset='utf-8'>
            <title>文档总结报告</title>
            <style>{self.css_styles}</style>
        </head>
        <body>
            <h1>📑 文档总结报告</h1>
            <h2>总体摘要</h2>
            <div class="summary">{self.final_summary}</div>
            {self.format_failed_files()}
            <div class="details">
                <h2>📚 各文件详细总结</h2>
                {self.format_file_summaries()}
            </div>
        </body>
        </html>
        """


